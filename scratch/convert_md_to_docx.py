import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []
    in_code_block = False
    code_content = []

    for line in lines:
        stripped = line.strip()
        
        # Code Block Toggle
        if stripped.startswith('```'):
            if in_code_block:
                # Close code block
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                in_code_block = False
                code_content = []
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line.rstrip())
            continue

        # Table detection
        if '|' in line and not in_code_block:
            if not in_table:
                in_table = True
                table_data = []
            
            # Skip separator line like | :--- | :--- |
            if ':---' in line or '---:' in line:
                continue
                
            cells = [c.strip() for c in line.split('|') if c.strip() or (c == '' and line.count('|') > 1)]
            # Filter out empty first/last cells if they are just artifacts of the md format | cell |
            if line.strip().startswith('|'):
                cells = [c.strip() for c in line.split('|')][1:-1]
            else:
                cells = [c.strip() for c in line.split('|')]
            
            if cells:
                table_data.append(cells)
            continue
        else:
            if in_table:
                # Create the table in docx
                if table_data:
                    rows = len(table_data)
                    cols = len(table_data[0])
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for r in range(rows):
                        for c in range(min(cols, len(table_data[r]))):
                            table.cell(r, c).text = table_data[r][c]
                in_table = False
                table_data = []

        # Headings
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        # Lists
        elif stripped.startswith('* ') or stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        elif stripped.startswith('1. ') or stripped.startswith('2. '):
             # Simplified numbered list detection
             doc.add_paragraph(stripped[3:], style='List Number')
        # Blockquotes (simplified)
        elif stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            p.style = 'Quote'
        # Empty line
        elif not stripped:
            continue
        # Horizontal rule
        elif stripped == '---':
            doc.add_page_break() # Or just skip
        # Normal text
        else:
            doc.add_paragraph(stripped)

    # Save
    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("Usage: python convert_md_to_docx.py <input.md> <output.docx>")
    else:
        md_file = sys.argv[1]
        docx_file = sys.argv[2]
        convert_md_to_docx(md_file, docx_file)
